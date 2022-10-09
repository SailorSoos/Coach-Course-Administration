#Welcome! You are executing this to send a single certificate off for the first time.
import os
import docx
import re
from docx.api import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import win32com.client
import os.path

course_participant = input("Participant name:")
participant_email = input("Email:")
print("Cert type from list below: \nBuddy = b \nAssistant = a \nLTS =  l \nHead = h \nRace = ra \nKeelboat 1 = k1 \nKeelboat 2 = k2 \nKeelboat 3 = k3")
print("\n")
certification_type = input("Cert type: ")
print("")
#Update this date yearly
expiry_date = "30 June, 2026"

#Email signature details
line_1 = ('<b><font color="rgb(0,65,92)"> Peter Soosalu | Coach Development Manager | Yachting New Zealand </font></b> <br>')
line_2 = ('<b><font color="rgb(0,65,92)">M</b></font> <font color="rgb(0,65,92)">(021) 037 2419 </font>| <b><font color="rgb(0,65,92)">E</font></b> peters@yachting.org.nz <br>')
line_3 = ('Yachtingnz.org.nz | <a href="https://www.facebook.com/YachtingNewZealand/">Facebook</a> | <a href="https://www.facebook.com/NZLSailingTeam/">NZL Sailing Team</a>') 
line_4 = ('<br><br> For the latest news and offerings download the Yachting New Zealand app.') 
line_5 = ('<br><a href="https://apps.apple.com/us/app/yachting-nz-app/id1040333130?amp%3Bls=1&amp%3Bamp%3Bmt=8&amp%3Bl=nl"><img src="https://developer.apple.com/news/images/download-on-the-app-store-badge.png" width="108" height="36" alt="app_store"></a> <a href="https://play.google.com/store/apps/details?id=com.app.p1107GC"><img src="https://encrypted-tbn0.gstatic.com/images?q=tbn:ANd9GcT_E6ZM5CF_cPm4tzqW6MpFGm2efBY1QL6v6w&usqp=CAU" width="108" height="36" alt="android_store"></a>')
line_6 = ('<br><br><img src="https://i.ibb.co/w7Lfz35/99798ce7-9981-4308-8295-3b3fa5386314.jpg" alt="ynz_banner">')
line_7 = ('<br><br> <font size="-2" color="rgb(220,220,220)"> The content of this e-mail is confidential and may contain copyright information. If you are not the intended recipient, please delete the </font><br> <font size="-2" color="rgb(220,220,220)">message and notify the sender immediately. You should scan this message and any attached files for viruses. We accept no liability for </font><br><font size="-2" color="rgb(220,220,220)"> any loss caused either directly or indirectly by a virus arising from the use of this message or any attached file. Thank you. </font>')
email_signature = (line_1 + line_2 + line_3 + line_4 + line_5 + line_6 + line_7)

#Open the certificate template
if certification_type == 'a':
        doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\LTSAssistant.docx')
elif certification_type == 'l':
        doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\LTS.docx')
elif certification_type == 'b':
        doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\LTSBuddy.docx')
elif certification_type == 'h':
        doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\LTSHead.docx')
elif certification_type == 'ra':
        doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\Race.docx')
elif certification_type == 'k1':
        doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\K1.docx')
elif certification_type == 'k2':
        doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\K2.docx')
elif certification_type == 'k3':
        doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\K3.docx')

#Change the font size and type
obj_styles = doc.styles
obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.PARAGRAPH)
obj_font = obj_charstyle.font
obj_font.size = Pt(80)
obj_font.name = 'Freestyle Script'

#Update course_participant
for paragraph in doc.paragraphs:
        if 'name_here' in paragraph.text:
                paragraph.style = doc.styles['CommentsStyle']
                paragraph.text = course_participant
        
#Where to save it, what to save it as in .docx
if certification_type == 'a':
        doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Assistant, ') + course_participant + ('.docx'))
elif certification_type == 'l':
        doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS, ') + course_participant + ('.docx'))
elif certification_type == 'b':
        doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Buddy, ') + course_participant + ('.docx'))
elif certification_type == 'h':
        doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Head Coach, ') + course_participant + ('.docx'))
elif certification_type == 'ra':
        doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Race Coach, ') + course_participant + ('.docx'))
elif certification_type == 'k1':
        doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 1, ') + course_participant + ('.docx'))
elif certification_type == 'k2':
        doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 2, ') + course_participant + ('.docx'))
elif certification_type == 'k3':
        doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 3, ') + course_participant + ('.docx'))

#Convert word doc to pdf
from docx2pdf import convert
if certification_type == 'a':
        convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Assistant, ') + course_participant + ('.docx'))
elif certification_type == 'l':
        convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS, ') + course_participant + ('.docx'))
elif certification_type == 'b':
        convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Buddy, ') + course_participant + ('.docx'))
elif certification_type == 'h':
        convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Head Coach, ') + course_participant + ('.docx'))
elif certification_type == 'ra':
        convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Race Coach, ') + course_participant + ('.docx'))
elif certification_type == 'k1':
        convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 1, ') + course_participant + ('.docx'))
elif certification_type == 'k2':
        convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 2, ') + course_participant + ('.docx'))
elif certification_type == 'k3':
        convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 3, ') + course_participant + ('.docx'))

#Delete the word document version
if certification_type == 'a':
        os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Assistant, ') + course_participant + ('.docx'))
elif certification_type == 'l':
        os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS, ') + course_participant + ('.docx'))
elif certification_type == 'b':
        os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Buddy, ') + course_participant + ('.docx'))
elif certification_type == 'h':
        os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Head Coach, ') + course_participant + ('.docx'))
elif certification_type == 'ra':
        os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Race Coach, ') + course_participant + ('.docx'))
elif certification_type == 'k1':
        os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 1, ') + course_participant + ('.docx'))
elif certification_type == 'k2':
        os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 2, ') + course_participant + ('.docx'))
elif certification_type == 'k3':
        os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 3, ') + course_participant + ('.docx'))


#Outlook mail portion
outlook = win32com.client.Dispatch('outlook.application')
mail = outlook.CreateItem(0)

#Need to customize email for the qualification
if certification_type == 'a':
        email = ("Dear ") + course_participant + (",<br><br>You have successfully completed the Yachting New Zealand Learn to Sail Coach Course. I am pleased to advise that you are now an officially recognised <b> Assistant Learn to Sail Coach. </b>"
                        "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                        "As an Assistant Learn to Sail Coach you are qualified to assist with the Yachting New Zealand Learn to Sail Dinghy (Level I and II) program. These levels can be taught at yacht clubs and other Yachting New Zealand affiliated organisations subject to the Yachting New Zealand safety requirements. <br><br>"
                        "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b>Assistant Learn to Sail Coach (Dinghy) Certificate.</b> You can upgrade to a Learn to sail coach when you turn 18 or working with the feedback the coach developer has given to upgrade your certificate. When you are ready to upgrade, you can by filling out a revalidation from, available to download off the Yachting New Zealand website.   <br><br>"
                        "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                        "If you are interested in furthering your coaching experience you may be keen to look at some becoming a Race Coach – information can be found under the <b>Coaches</b> page on the Yachting New Zealand website. Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                        "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                        "Regards,<br><br>") + email_signature
elif certification_type == 'l':
        email = ("Dear ") + course_participant + (",<br><br>You have successfully completed the Yachting New Zealand Learn to Sail Coach Course. I am pleased to advise that you are now an officially recognised <b> Learn to Sail Coach. </b>"
                        "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                        "As a Learn to Sail Coach you are qualified to teach all aspects of the Yachting New Zealand Learn to Sail Dinghy (Level I and II) program. These levels can be taught at yacht clubs and other Yachting New Zealand affiliated organisations subject to the Yachting New Zealand safety requirements. <br><br>"
                        "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b> Learn to Sail Coach (Dinghy) Certificate.</b> You can upgrade to a Learn to sail coach when you turn 18 or working with the feedback the coach developer has given to upgrade your certificate. When you are ready to upgrade, you can by filling out a revalidation from, available to download off the Yachting New Zealand website.   <br><br>"
                        "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                        "If you are interested in furthering your coaching experience you may be keen to look at some becoming a Race Coach – information can be found under the <b>Coaches</b> page on the Yachting New Zealand website. Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                        "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                        "Regards,<br><br>") + email_signature
elif certification_type == 'b':
        email = ("Dear ") + course_participant + (",<br><br>You have successfully completed the Yachting New Zealand Learn to Sail Coach Course. I am pleased to advise that you are now an officially recognised <b> Buddy Learn to Sail Coach. </b>"
                        "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                        "As a Buddy Learn to Sail Coach you are qualified to assist with the Yachting New Zealand Learn to Sail Dinghy (Level I and II) program. These levels can be taught at yacht clubs and other Yachting New Zealand affiliated organisations subject to the Yachting New Zealand safety requirements. <br><br>"
                        "A buddy coach must always be working under the supervision of a fully qualified Learn to Sail coach. <br><br>"
                        "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b>Buddy Learn to Sail Coach (Dinghy) Certificate.</b> You can upgrade to an Assistant Learn to sail coach when you turn 15 or working with the feedback the coach developer has given to upgrade your certificate. When you are ready to upgrade, you can by filling out a revalidation from, available to download off the Yachting New Zealand website.   <br><br>"
                        "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                        "If you are interested in furthering your coaching experience you may be keen to look at some becoming a Race Coach – information can be found under the <b>Coaches</b> page on the Yachting New Zealand website. Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                        "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                        "Regards,<br><br>") + email_signature
elif certification_type == 'h':
        email = ("Dear ") + course_participant + (",<br><br>You have successfully completed the Yachting New Zealand Learn to Sail Coach Course. I am pleased to advise that you are now an officially recognised <b> Head Learn to Sail Coach. </b>"
                        "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                        "As a Head Learn to Sail Coach you are qualified to teach all aspects of the Yachting New Zealand Learn to Sail Dinghy (Level I and II) program. These levels can be taught at yacht clubs and other Yachting New Zealand affiliated organisations subject to the Yachting New Zealand safety requirements. Additionally, as a Head Learn to Sail coach, mentoring and supporting other coaches in your area is not only encouraged, but can help improve your own coaching experience by sharing ideas and seeing new ones. <br><br>"
                        "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b> Head Learn to Sail Coach (Dinghy) Certificate.</b> <br><br>"
                        "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                        "If you are interested in furthering your coaching experience you may be keen to look at some becoming a Race Coach – information can be found under the <b>Coaches</b> page on the Yachting New Zealand website. Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                        "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                        "Regards,<br><br>") + email_signature
elif certification_type == 'ra':
        email = ("Dear ") + course_participant + (",<br><br>I am pleased to advise that you are now an officially recognised <b> Race Coach. </b>"
                        "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                        "The Race Coach is the first step along the race coach pathway, followed by the Regatta coach, Performance coach, and finally Olympic coach. As a race coach, mentoring and supporting other coaches in your area is not only encouraged, but can help improve your own coaching experience by sharing ideas and seeing new ones. <br><br>"
                        "Your qualification does not have a set required sailor to coach ratio, but it is recommended to have a max ratio of 6:1.<br><br>"
                        "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b> Race Coach Certificate.</b> <br><br>"
                        "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                        "Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                        "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                        "Regards,<br><br>") + email_signature
elif certification_type == 'k1':
        email = ("Dear ") + course_participant + (",<br><br>I am pleased to advise that you are now an officially recognised <b> Keelboat Level 1 Coach. </b>"
                        "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                        "As a keelboat coach, mentoring and supporting other coaches in your area is not only encouraged, but can help improve your own coaching experience by sharing ideas and seeing new ones. <br><br>"
                        "Your qualification does not have a set required sailor to coach ratio, but it is recommended to have a max ratio of 7:1.<br><br>"
                        "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b> Keelboat level 1 Coach Certificate.</b> <br><br>"
                        "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                        "Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                        "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                        "Regards,<br><br>") + email_signature
elif certification_type == 'k2':
        email = ("Dear ") + course_participant + (",<br><br>I am pleased to advise that you are now an officially recognised <b> Keelboat Level 2 Coach. </b>"
                        "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                        "As a keelboat coach, mentoring and supporting other coaches in your area is not only encouraged, but can help improve your own coaching experience by sharing ideas and seeing new ones. <br><br>"
                        "Your qualification does not have a set required sailor to coach ratio, but it is recommended to have a max ratio of 7:1.<br><br>"
                        "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b> Keelboat level 2 Certificate.</b> <br><br>"
                        "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                        "Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                        "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                        "Regards,<br><br>") + email_signature
elif certification_type == 'k3':
        email = ("Dear ") + course_participant + (",<br><br>I am pleased to advise that you are now an officially recognised <b> Keelboat Level 3 Coach. </b>"
                        "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                        "As a keelboat coach, mentoring and supporting other coaches in your area is not only encouraged, but can help improve your own coaching experience by sharing ideas and seeing new ones. <br><br>"
                        "Your qualification does not have a set required sailor to coach ratio, but it is recommended to have a max ratio of 7:1.<br><br>"
                        "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b> Keelboat level 3 Certificate.</b> <br><br>"
                        "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                        "Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                        "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                        "Regards,<br><br>") + email_signature

#Create the email 
mail.To = participant_email
mail.Subject = 'Yachting New Zealand - Coaching Course Certificate - ' + course_participant 
mail.HTMLBody = email

# Certification attachments
if certification_type == 'a':
        course_certificate =  (('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Assistant, ') + course_participant + ('.pdf'))
elif certification_type == 'l':
        course_certificate =  (('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS, ') + course_participant + ('.pdf'))
elif certification_type == 'b':
        course_certificate =  (('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Buddy, ') + course_participant + ('.pdf'))
elif certification_type == 'h':
        course_certificate =  (('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Head Coach, ') + course_participant + ('.pdf'))
elif certification_type == 'ra':
        course_certificate =  (('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Race Coach, ') + course_participant + ('.pdf'))
elif certification_type == 'k1':
        course_certificate =  (('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 1, ') + course_participant + ('.pdf'))
elif certification_type == 'k2':
        course_certificate =  (('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 2, ') + course_participant + ('.pdf'))
elif certification_type == 'k3':
        course_certificate =  (('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 3, ') + course_participant + ('.pdf'))
mail.attachments.Add(course_certificate)
if certification_type == 'a':
        mail.attachments.Add('c:\\Users\Peters\Documents\CDM\Certificates\Attachments\YACHTING NEW ZEALAND SAFETY REQUIREMENT.pdf')
elif certification_type == 'l':
        mail.attachments.Add('c:\\Users\Peters\Documents\CDM\Certificates\Attachments\YACHTING NEW ZEALAND SAFETY REQUIREMENT.pdf')
elif certification_type == 'b':
        mail.attachments.Add('c:\\Users\Peters\Documents\CDM\Certificates\Attachments\YACHTING NEW ZEALAND SAFETY REQUIREMENT.pdf')
elif certification_type == 'h':
        mail.attachments.Add('c:\\Users\Peters\Documents\CDM\Certificates\Attachments\YACHTING NEW ZEALAND SAFETY REQUIREMENT.pdf')
mail.Send()

print("")
print(('Certification email sent to: ') + course_participant)

#cc another person (alt email from crm?)
#mail.CC = 'peters@yachtingnz.org.nz'



#Disclaimer:
#Moving this program, or any files associated with it, will cause it to break. Be sure to check and update the following if you move this
#Install the necessary pip extensions
#Move the default certificate templates and update root directories
#Specify where you want the auto created certs dropped, update root directories
#Update where Ynz safety req's goes, update root directories