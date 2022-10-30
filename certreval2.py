#Welcome! You are executing this to revalidate a certificate.
course_participant = input("Participant name:")
print("")
participant_email = input("Email:")
print("\nCert type from list below:")
print("'assistcert' \n'ltscert' \n'headcert' \n'racecert' \n'keel1cert' \n'keel2cert' \n'keel2cert' \n'keel3cert' \n")
certification_type = input("Cert type: ")
#Update this date yearly
expiry_date = "30 June, 2025"

import os
import docx
import re
from docx.api import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

#Email signature details
email_signature = ('<b><font color="rgb(0,65,92)"> Peter Soosalu | Coach Development Manager | Yachting New Zealand </font></b> <br>'
        '<b><font color="rgb(0,65,92)">M</b></font> <font color="rgb(0,65,92)">(021) 037 2419 </font>| <b><font color="rgb(0,65,92)">E</font></b> peters@yachting.org.nz <br>'
        '<a href="http://www.yachtingnz.org.nz">Yachtingnz.org.nz</a> | <a href="https://www.facebook.com/YachtingNewZealand/">Facebook</a> | <a href="https://www.facebook.com/NZLSailingTeam/">NZL Sailing Team</a>'
        '<br><br> For the latest news and offerings download the Yachting New Zealand app.'
        '<br><a href="https://apps.apple.com/us/app/yachting-nz-app/id1040333130?amp%3Bls=1&amp%3Bamp%3Bmt=8&amp%3Bl=nl"><img src="https://developer.apple.com/news/images/download-on-the-app-store-badge.png" width="108" height="36" alt="app_store"></a> <a href="https://play.google.com/store/apps/details?id=com.app.p1107GC"><img src="https://encrypted-tbn0.gstatic.com/images?q=tbn:ANd9GcT_E6ZM5CF_cPm4tzqW6MpFGm2efBY1QL6v6w&usqp=CAU" width="108" height="36" alt="android_store"></a>'
        '<br><br><img src="https://i.ibb.co/w7Lfz35/99798ce7-9981-4308-8295-3b3fa5386314.jpg" alt="ynz_banner">'
        '<br><br> <font size="-2" color="rgb(220,220,220)"> The content of this e-mail is confidential and may contain copyright information. If you are not the intended recipient, please delete the </font><br> <font size="-2" color="rgb(220,220,220)">message and notify the sender immediately. You should scan this message and any attached files for viruses. We accept no liability for </font><br><font size="-2" color="rgb(220,220,220)"> any loss caused either directly or indirectly by a virus arising from the use of this message or any attached file. Thank you. </font>')


#Need to customize email for the qualification
if certification_type == 'assistcert':
        email = ("Dear ") + course_participant + (",<br><br>After reviewing your revalidation request, I am happy to revalidate your Yachting New Zealand Learn to Sail coach certificate. You are now an officially recognised <b> Assistant Learn to Sail Coach. </b>"
                        "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                        "As an Assistant Learn to Sail Coach you are qualified to assist with the Yachting New Zealand Learn to Sail Dinghy (Level I and II) program. These levels can be taught at yacht clubs and other Yachting New Zealand affiliated organisations subject to the Yachting New Zealand safety requirements. <br><br>"
                        "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b>Assistant Learn to Sail Coach (Dinghy) Certificate.</b> You can upgrade to a Learn to sail coach when you turn 18 or working with the feedback the coach developer has given to upgrade your certificate. When you are ready to upgrade, you can by filling out a revalidation from, available to download off the Yachting New Zealand website.   <br><br>"
                        "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                        "If you are interested in furthering your coaching experience you may be keen to look at some becoming a Race Coach – information can be found under the <b>Coaches</b> page on the Yachting New Zealand website. Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                        "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                        "Regards,<br><br>") + email_signature
elif certification_type == 'ltscert':
        email = ("Dear ") + course_participant + (",<br><br>After reviewing your revalidation request, I am happy to revalidate your Yachting New Zealand Learn to Sail coach certificate. You are now an officially recognised <b> Learn to Sail Coach. </b>"
                        "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                        "As a Learn to Sail Coach you are qualified to teach all aspects of the Yachting New Zealand Learn to Sail Dinghy (Level I and II) program. These levels can be taught at yacht clubs and other Yachting New Zealand affiliated organisations subject to the Yachting New Zealand safety requirements. <br><br>"
                        "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b> Learn to Sail Coach (Dinghy) Certificate.</b> You can upgrade to a Learn to sail coach when you turn 18 or working with the feedback the coach developer has given to upgrade your certificate. When you are ready to upgrade, you can by filling out a revalidation from, available to download off the Yachting New Zealand website.   <br><br>"
                        "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                        "If you are interested in furthering your coaching experience you may be keen to look at some becoming a Race Coach – information can be found under the <b>Coaches</b> page on the Yachting New Zealand website. Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                        "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                        "Regards,<br><br>") + email_signature
elif certification_type == 'headcert':
        email = ("Dear ") + course_participant + (",<br><br>After reviewing your revalidation request, I am happy to revalidate your Yachting New Zealand Learn to Sail coach certificate. You are now an officially recognised <b> Head Learn to Sail Coach. </b>"
                        "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                        "As a Head Learn to Sail Coach you are qualified to teach all aspects of the Yachting New Zealand Learn to Sail Dinghy (Level I and II) program. These levels can be taught at yacht clubs and other Yachting New Zealand affiliated organisations subject to the Yachting New Zealand safety requirements. Additionally, as a Head Learn to Sail coach, mentoring and supporting other coaches in your area is not only encouraged, but can help improve your own coaching experience by sharing ideas and seeing new ones. <br><br>"
                        "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b> Head Learn to Sail Coach (Dinghy) Certificate.</b> <br><br>"
                        "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                        "If you are interested in furthering your coaching experience you may be keen to look at some becoming a Race Coach – information can be found under the <b>Coaches</b> page on the Yachting New Zealand website. Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                        "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                        "Regards,<br><br>") + email_signature
elif certification_type == 'racecert':
        email = ("Dear ") + course_participant + (",<br><br>After reviewing your revalidation request, I am happy to revalidate your Yachting New Zealand Race coach certificate. You are now an officially recognised <b> Race Coach. </b>"
                        "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                        "Additionally, as a Race coach, mentoring and supporting other coaches in your area is not only encouraged, but can help improve your own coaching experience by sharing ideas and seeing new ones. <br><br>"
                        "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b> Race Coach Certificate.</b> <br><br>"
                        "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                        "If you are interested in furthering your coaching experience you may be keen to look at some becoming a Regatta Coach – information can be found under the <b>Coaches</b> page on the Yachting New Zealand website. Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                        "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                        "Regards,<br><br>") + email_signature
elif certification_type == 'keel1cert':
        email = ("Dear ") + course_participant + (",<br><br>After reviewing your revalidation request, I am happy to revalidate your Yachting New Zealand Keelboat Level 1 coach certificate. You are now an officially recognised <b> Level 1 Keelboat Coach. </b>"
                        "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                        "Additionally, as a Keelboat coach, mentoring and supporting other coaches in your area is not only encouraged, but can help improve your own coaching experience by sharing ideas and seeing new ones. <br><br>"
                        "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b> Keelboat Level 1 Coach Certificate.</b> <br><br>"
                        "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                        "Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                        "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                        "Regards,<br><br>") + email_signature
elif certification_type == 'keel2cert':
        email = ("Dear ") + course_participant + (",<br><br>After reviewing your revalidation request, I am happy to revalidate your Yachting New Zealand Keelboat Level 2 coach certificate. You are now an officially recognised <b> Level 2 Keelboat Coach. </b>"
                        "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                        "Additionally, as a Keelboat coach, mentoring and supporting other coaches in your area is not only encouraged, but can help improve your own coaching experience by sharing ideas and seeing new ones. <br><br>"
                        "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b> Keelboat Level 2 Coach Certificate.</b> <br><br>"
                        "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                        "Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                        "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                        "Regards,<br><br>") + email_signature
elif certification_type == 'keel3cert':
        email = ("Dear ") + course_participant + (",<br><br>After reviewing your revalidation request, I am happy to revalidate your Yachting New Zealand Keelboat Level 3 coach certificate. You are now an officially recognised <b> Level 3 Keelboat Coach. </b>"
                        "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                        "Additionally, as a Keelboat coach, mentoring and supporting other coaches in your area is not only encouraged, but can help improve your own coaching experience by sharing ideas and seeing new ones. <br><br>"
                        "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b> Keelboat Level 3 Coach Certificate.</b> <br><br>"
                        "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                        "Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                        "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                        "Regards,<br><br>") + email_signature

#Open the certificate template
if certification_type == 'assistcert':
        doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\LTSAssistant.docx')
elif certification_type == 'ltscert':
        doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\LTS.docx')
elif certification_type == 'buddycert':
        doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\LTSBuddy.docx')
elif certification_type == 'headcert':
        doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\LTSHead.docx')
elif certification_type == 'racecert':
        doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\Race.docx')
elif certification_type == 'keel1cert':
        doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\K1.docx')
elif certification_type == 'keel2cert':
        doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\K2.docx')
elif certification_type == 'keel3cert':
        doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\K3.docx')

#Change the font size and type
obj_styles = doc.styles
obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.PARAGRAPH)
obj_font = obj_charstyle.font
obj_font.size = Pt(80)
obj_font.name = 'Freestyle Script'

#Update course_participant
for paragraph in doc.paragraphs:
        if 'name_here' in paragraph.text:
                paragraph.style = doc.styles['CommentsStyle']
                paragraph.text = course_participant
        
#Where to save it, what to save it as in .docx
if certification_type == 'assistcert':
        doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Assistant, ') + course_participant + ('.docx'))
elif certification_type == 'ltscert':
        doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS, ') + course_participant + ('.docx'))
elif certification_type == 'headcert':
        doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Head Coach, ') + course_participant + ('.docx'))
elif certification_type == 'racecert':
        doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Race Coach, ') + course_participant + ('.docx'))
elif certification_type == 'keel1cert':
        doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 1, ') + course_participant + ('.docx'))
elif certification_type == 'keel2cert':
        doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 2, ') + course_participant + ('.docx'))
elif certification_type == 'keel3cert':
        doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 3, ') + course_participant + ('.docx'))

#Convert word doc to pdf
from docx2pdf import convert
if certification_type == 'assistcert':
        convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Assistant, ') + course_participant + ('.docx'))
elif certification_type == 'ltscert':
        convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS, ') + course_participant + ('.docx'))
elif certification_type == 'headcert':
        convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Head Coach, ') + course_participant + ('.docx'))
elif certification_type == 'racecert':
        convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Race Coach, ') + course_participant + ('.docx'))
elif certification_type == 'keel1cert':
        convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 1, ') + course_participant + ('.docx'))
elif certification_type == 'keel2cert':
        convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 2, ') + course_participant + ('.docx'))
elif certification_type == 'keel3cert':
        convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 3, ') + course_participant + ('.docx'))

#Delete the word document version
if certification_type == 'assistcert':
        os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Assistant, ') + course_participant + ('.docx'))
elif certification_type == 'ltscert':
        os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS, ') + course_participant + ('.docx'))
elif certification_type == 'headcert':
        os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Head Coach, ') + course_participant + ('.docx'))
elif certification_type == 'racecert':
        os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Race Coach, ') + course_participant + ('.docx'))
elif certification_type == 'keel1cert':
        os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 1, ') + course_participant + ('.docx'))
elif certification_type == 'keel2cert':
        os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 2, ') + course_participant + ('.docx'))
elif certification_type == 'keel3cert':
        os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 3, ') + course_participant + ('.docx'))


#Outlook mail portion
import win32com.client
import os.path
outlook = win32com.client.Dispatch('outlook.application')
mail = outlook.CreateItem(0)

#Create the email 
mail.To = participant_email
mail.Subject = 'Yachting New Zealand - Coaching Course Certificate - ' + course_participant 
mail.HTMLBody = email

# Certification attachments
if certification_type == 'assistcert':
        course_certificate =  (('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Assistant, ') + course_participant + ('.pdf'))
elif certification_type == 'ltscert':
        course_certificate =  (('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS, ') + course_participant + ('.pdf'))
elif certification_type == 'headcert':
        course_certificate =  (('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Head Coach, ') + course_participant + ('.pdf'))
elif certification_type == 'racecert':
        course_certificate =  (('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Race Coach, ') + course_participant + ('.pdf'))
elif certification_type == 'keel1cert':
        course_certificate =  (('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 1, ') + course_participant + ('.pdf'))
elif certification_type == 'keel2cert':
        course_certificate =  (('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 2, ') + course_participant + ('.pdf'))
elif certification_type == 'keel3cert':
        course_certificate =  (('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 3, ') + course_participant + ('.pdf'))
mail.attachments.Add(course_certificate)
mail.attachments.Add('c:\\Users\Peters\Documents\CDM\Certificates\Attachments\YACHTING NEW ZEALAND SAFETY REQUIREMENT.pdf')
mail.Send()

print("")
print(('Revalidation email sent to: ') + course_participant)

#cc another person (alt email from crm?)
#mail.CC = 'peters@yachtingnz.org.nz'