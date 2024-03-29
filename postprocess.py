from numpy import NaN
import pandas as pd
import os
import docx
import re
#from sqlalchemy import null
import win32com.client
import os.path
from docx.api import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx2pdf import convert
excel_sheet = pd.read_excel("c:/Users/Peters/Downloads/process_these.xlsx")

email_signature = ('<b><font color="rgb(0,65,92)"> Peter Soosalu | Coach Development Manager | Yachting New Zealand </font></b> <br>'
        '<b><font color="rgb(0,65,92)">M</b></font> <font color="rgb(0,65,92)">(021) 037 2419 </font>| <b><font color="rgb(0,65,92)">E</font></b> peters@yachting.org.nz <br>'
        '<a href="http://www.yachtingnz.org.nz">Yachtingnz.org.nz</a> | <a href="https://www.facebook.com/YachtingNewZealand/">Facebook</a> | <a href="https://www.facebook.com/NZLSailingTeam/">NZL Sailing Team</a>'
        '<br><br> For the latest news and offerings download the Yachting New Zealand app.'
        '<br><a href="https://apps.apple.com/us/app/yachting-nz-app/id1040333130?amp%3Bls=1&amp%3Bamp%3Bmt=8&amp%3Bl=nl"><img src="https://developer.apple.com/news/images/download-on-the-app-store-badge.png" width="108" height="36" alt="app_store"></a> <a href="https://play.google.com/store/apps/details?id=com.app.p1107GC"><img src="https://encrypted-tbn0.gstatic.com/images?q=tbn:ANd9GcT_E6ZM5CF_cPm4tzqW6MpFGm2efBY1QL6v6w&usqp=CAU" width="108" height="36" alt="android_store"></a>'
        '<br><br><img src="https://i.ibb.co/w7Lfz35/99798ce7-9981-4308-8295-3b3fa5386314.jpg" alt="ynz_banner">'
        '<br><br> <font size="-2" color="rgb(220,220,220)"> The content of this e-mail is confidential and may contain copyright information. If you are not the intended recipient, please delete the </font><br> <font size="-2" color="rgb(220,220,220)">message and notify the sender immediately. You should scan this message and any attached files for viruses. We accept no liability for </font><br><font size="-2" color="rgb(220,220,220)"> any loss caused either directly or indirectly by a virus arising from the use of this message or any attached file. Thank you. </font>')

#refine these into effective functions
#email function
def email(course_participant, course_certificate, participant_email, course_type, path):
        
        LTSbody = (("Dear ") + course_participant + (",<br><br>You have successfully completed the Yachting New Zealand Learn to Sail Coach Course. I am pleased to advise that you are now an officially recognised <b> Assistant Learn to Sail Coach. </b>"
                "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                "As a ") + course_type1 + (" you are qualified to assist with the Yachting New Zealand Learn to Sail Dinghy (Level I and II) program. These levels can be taught at yacht clubs and other Yachting New Zealand affiliated organisations subject to the Yachting New Zealand safety requirements. <br><br>"
                "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b>") + course_type2 + ("</b> You can upgrade to a Learn to sail coach when you turn 18 or working with the feedback the coach developer has given to upgrade your certificate. When you are ready to upgrade, you can by filling out a revalidation from, available to download off the Yachting New Zealand website.   <br><br>"
                "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                "If you are interested in furthering your coaching experience you may be keen to look at some becoming a Race Coach – information can be found under the <b>Coaches</b> page on the Yachting New Zealand website. Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                "Regards,<br><br>") + email_signature)

        KBbody = (("Dear ") + course_participant + (",<br><br>I am pleased to advise that you are now an officially recognised <b>") + registrations_qual + (". </b>"
                "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                "As a keelboat coach, mentoring and supporting other coaches in your area is not only encouraged, but can help improve your own coaching experience by sharing ideas and seeing new ones. <br><br>"
                "Your qualification does not have a set required sailor to coach ratio, but it is recommended to have a max ratio of 7:1.<br><br>"
                "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b> ") + registrations_qual + (" Certificate.</b> <br><br>"
                "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                "Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                "Regards,<br><br>") + email_signature)

        Racebody = (("Dear ") + course_participant + (",<br><br>I am pleased to advise that you are now an officially recognised <b> Race Coach. </b>"
                "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                "The Race Coach is the first step along the race coach pathway, followed by the Regatta coach, Performance coach, and finally Olympic coach. As a race coach, mentoring and supporting other coaches in your area is not only encouraged, but can help improve your own coaching experience by sharing ideas and seeing new ones. <br><br>"
                "Your qualification does not have a set required sailor to coach ratio, but it is recommended to have a max ratio of 6:1.<br><br>"
                "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b> Race Coach Certificate.</b> <br><br>"
                "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                "Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                "Regards,<br><br>") + email_signature)
        
        outlook = win32com.client.Dispatch('outlook.application')
        mail = outlook.CreateItem(0)
        mail.To = participant_email
        mail.Subject = 'Yachting New Zealand - Coaching Course Certificate - ' + course_participant 
        mail.HTMLBody = LTSbody
        course_certificate =  path + ('\Auto_created_certs\Yachting New Zealand - LTS Assistant, ') + course_participant + ('.pdf')
        mail.attachments.Add(course_certificate)
        mail.attachments.Add('c:\\Users\Peters\Documents\CDM\Certificates\Attachments\YACHTING NEW ZEALAND SAFETY REQUIREMENT.pdf')
        mail.Send()

#make and store course cert function
def certificate(course_participant, course_type, template_path, output_path):
        doc = docx.Document(template_path + course_type + '.docx')
        obj_styles = doc.styles
        obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.PARAGRAPH)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(80)
        obj_font.name = 'Freestyle Script'

        for paragraph in doc.paragraphs:
                if 'name_here' in paragraph.text:
                        paragraph.style = doc.styles['CommentsStyle']
                        paragraph.text = course_participant

        doc.save((output_path + 'Yachting New Zealand - ') + course_participant + ('.docx'))
        convert((output_path + 'Yachting New Zealand - ') + course_participant + ('.docx'))
        os.remove((output_path + 'Yachting New Zealand - ') + course_participant + ('.docx'))

i = 0
while i < len(excel_sheet):
        registrations_name1 = excel_sheet.iloc[i, 3]
        registrations_name = str(registrations_name1)
        registrations_qual1 = excel_sheet.iloc[i, 7]
        registrations_qual = str(registrations_qual1)
        registrations_expiry1 = excel_sheet.iloc[i, 8]
        registrations_expiry =str(registrations_expiry1)
        registrations_email1 = excel_sheet.iloc[i, 11]
        registrations_email = str(registrations_email1)
        participant_details = [registrations_name, registrations_qual, registrations_expiry, registrations_email]

#Generic course related details trying to loop
        course_participant = participant_details[0]
        if registrations_qual == 'Assistant LTS Coach (Dinghy)':
                course_certification = 'assistcert'
        elif registrations_qual == 'LTS Coach (Dinghy)':
                course_certification = 'ltscert'
        elif registrations_qual == 'LTS Buddy':
                course_certification = 'buddycert'
        elif registrations_qual == 'Master LTS Coach (Dinghy)':
                course_certification = 'headcert'
        elif registrations_qual == 'LTS Coach (Keelboat) Level 1':
                course_certification = 'k1'
        elif registrations_qual == 'LTS Coach (Keelboat) Level 2':
                course_certification = 'k2'
        elif registrations_qual == 'LTS Coach (Keelboat) Level 3':
                course_certification = 'k3'
        else:
                course_certification = 'doembark'
        expiry_date = "30 June, 2025"
        participant_email = participant_details[3]

        #add this as a universal item in a class 
        path = ('c:\\Users\Peters\Documents\CDM\Certificates')
        course_type1 = registrations_qual
        course_type2 = "Assistant Learn to Sail Coach (Dinghy) Certificate."
        LTSbody = (("Dear ") + course_participant + (",<br><br>You have successfully completed the Yachting New Zealand Learn to Sail Coach Course. I am pleased to advise that you are now an officially recognised <b> Assistant Learn to Sail Coach. </b>"
                "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                "As a ") + course_type1 + (" you are qualified to assist with the Yachting New Zealand Learn to Sail Dinghy (Level I and II) program. These levels can be taught at yacht clubs and other Yachting New Zealand affiliated organisations subject to the Yachting New Zealand safety requirements. <br><br>"
                "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b>") + course_type2 + ("</b> You can upgrade to a Learn to sail coach when you turn 18 or working with the feedback the coach developer has given to upgrade your certificate. When you are ready to upgrade, you can by filling out a revalidation from, available to download off the Yachting New Zealand website.   <br><br>"
                "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                "If you are interested in furthering your coaching experience you may be keen to look at some becoming a Race Coach – information can be found under the <b>Coaches</b> page on the Yachting New Zealand website. Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                "Regards,<br><br>") + email_signature)

        KBbody = (("Dear ") + course_participant + (",<br><br>I am pleased to advise that you are now an officially recognised <b>") + registrations_qual + (". </b>"
                "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                "As a keelboat coach, mentoring and supporting other coaches in your area is not only encouraged, but can help improve your own coaching experience by sharing ideas and seeing new ones. <br><br>"
                "Your qualification does not have a set required sailor to coach ratio, but it is recommended to have a max ratio of 7:1.<br><br>"
                "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b> ") + registrations_qual + (" Certificate.</b> <br><br>"
                "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                "Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                "Regards,<br><br>") + email_signature)
        
#assistcert starts here
        if registrations_qual == 'Assistant LTS Coach (Dinghy)':

#Make a certificate for the course_participant
                doc = docx.Document(path + '\Auto_created_certs\Templates\LTSAssistant.docx')
                obj_styles = doc.styles
                obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.PARAGRAPH)
                obj_font = obj_charstyle.font
                obj_font.size = Pt(80)
                obj_font.name = 'Freestyle Script'

#Update course_participant
                for paragraph in doc.paragraphs:
                        if 'name_here' in paragraph.text:
                                paragraph.style = doc.styles['CommentsStyle']
                                paragraph.text = course_participant
        
#Save, convert to pdf, delete .docx
                doc.save((path + '\Auto_created_certs\Yachting New Zealand - LTS Assistant, ') + course_participant + ('.docx'))
                convert((path + '\Auto_created_certs\Yachting New Zealand - LTS Assistant, ') + course_participant + ('.docx'))
                os.remove((path + '\Auto_created_certs\Yachting New Zealand - LTS Assistant, ') + course_participant + ('.docx'))

#Outlook mail portion
                outlook = win32com.client.Dispatch('outlook.application')
                mail = outlook.CreateItem(0)
                mail.To = participant_email
                mail.Subject = 'Yachting New Zealand - Coaching Course Certificate - ' + course_participant 
                mail.HTMLBody = LTSbody

#Attachments
                course_certificate =  path + ('\Auto_created_certs\Yachting New Zealand - LTS Assistant, ') + course_participant + ('.pdf')
                mail.attachments.Add(course_certificate)
                mail.attachments.Add('c:\\Users\Peters\Documents\CDM\Certificates\Attachments\YACHTING NEW ZEALAND SAFETY REQUIREMENT.pdf')
                mail.Send()

                print(('Certification email sent to: ') + course_participant)
                i +=1


#ltscert starts here
        elif course_certification == 'ltscert':

#Make a certificate for the course_participant
                doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\LTS.docx')
                obj_styles = doc.styles
                obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.PARAGRAPH)
                obj_font = obj_charstyle.font
                obj_font.size = Pt(80)
                obj_font.name = 'Freestyle Script'

#Update course_participant
                for paragraph in doc.paragraphs:
                        if 'name_here' in paragraph.text:
                                paragraph.style = doc.styles['CommentsStyle']
                                paragraph.text = course_participant
        
#Save, convert to pdf, delete .docx
                doc.save(path + ('\Auto_created_certs\Yachting New Zealand - LTS, ') + course_participant + ('.docx'))
                convert(path + ('\Auto_created_certs\Yachting New Zealand - LTS, ') + course_participant + ('.docx'))
                os.remove(path + ('\Auto_created_certs\Yachting New Zealand - LTS, ') + course_participant + ('.docx'))

#Outlook mail portion
                outlook = win32com.client.Dispatch('outlook.application')
                mail = outlook.CreateItem(0)
                mail.To = participant_email
                mail.Subject = 'Yachting New Zealand - Coaching Course Certificate - ' + course_participant 
                mail.HTMLBody = ("Dear ") + course_participant + (",<br><br>You have successfully completed the Yachting New Zealand Learn to Sail Coach Course. I am pleased to advise that you are now an officially recognised <b> Learn to Sail Coach. </b>"
                        "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                        "As a Learn to Sail Coach you are qualified to teach all aspects of the Yachting New Zealand Learn to Sail Dinghy (Level I and II) program. These levels can be taught at yacht clubs and other Yachting New Zealand affiliated organisations subject to the Yachting New Zealand safety requirements. <br><br>"
                        "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b> Learn to Sail Coach (Dinghy) Certificate.</b> You can upgrade to a Learn to sail coach when you turn 18 or working with the feedback the coach developer has given to upgrade your certificate. When you are ready to upgrade, you can by filling out a revalidation from, available to download off the Yachting New Zealand website.   <br><br>"
                        "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                        "If you are interested in furthering your coaching experience you may be keen to look at some becoming a Race Coach – information can be found under the <b>Coaches</b> page on the Yachting New Zealand website. Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                        "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                        "Regards,<br><br>") + email_signature

#Attachments
                course_certificate =  ('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS, ') + course_participant + ('.pdf')
                mail.attachments.Add(course_certificate)
                mail.attachments.Add('c:\\Users\Peters\Documents\CDM\Certificates\Attachments\YACHTING NEW ZEALAND SAFETY REQUIREMENT.pdf')
                mail.Send()

                print(('Certification email sent to: ') + course_participant)
                i +=1



#headcert starts here
        elif course_certification == 'headcert':

#Make a certificate for the course_participant
                doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\LTShead.docx')
                obj_styles = doc.styles
                obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.PARAGRAPH)
                obj_font = obj_charstyle.font
                obj_font.size = Pt(80)
                obj_font.name = 'Freestyle Script'

#Update course_participant
                for paragraph in doc.paragraphs:
                        if 'name_here' in paragraph.text:
                                paragraph.style = doc.styles['CommentsStyle']
                                paragraph.text = course_participant
        
#Save, convert to pdf, delete .docx
                doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Head Coach, ') + course_participant + ('.docx'))
                convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Head Coach, ') + course_participant + ('.docx'))
                os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Head Coach, ') + course_participant + ('.docx'))

#Outlook mail portion
                outlook = win32com.client.Dispatch('outlook.application')
                mail = outlook.CreateItem(0)
                mail.To = participant_email
                mail.Subject = 'Yachting New Zealand - Coaching Course Certificate - ' + course_participant 
                mail.HTMLBody = ("Dear ") + course_participant + (",<br><br>You have successfully completed the Yachting New Zealand Learn to Sail Coach Course. I am pleased to advise that you are now an officially recognised <b> Head Learn to Sail Coach. </b>"
                        "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                        "As a Head Learn to Sail Coach you are qualified to teach all aspects of the Yachting New Zealand Learn to Sail Dinghy (Level I and II) program. These levels can be taught at yacht clubs and other Yachting New Zealand affiliated organisations subject to the Yachting New Zealand safety requirements. Additionally, as a Head Learn to Sail coach, mentoring and supporting other coaches in your area is not only encouraged, but can help improve your own coaching experience by sharing ideas and seeing new ones. <br><br>"
                        "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b> Head Learn to Sail Coach (Dinghy) Certificate.</b> <br><br>"
                        "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                        "If you are interested in furthering your coaching experience you may be keen to look at some becoming a Race Coach – information can be found under the <b>Coaches</b> page on the Yachting New Zealand website. Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                        "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                        "Regards,<br><br>") + email_signature

#Attachments
                course_certificate =  ('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Head Coach, ') + course_participant + ('.pdf')
                mail.attachments.Add(course_certificate)
                mail.attachments.Add('c:\\Users\Peters\Documents\CDM\Certificates\Attachments\YACHTING NEW ZEALAND SAFETY REQUIREMENT.pdf')
                mail.Send()

                print(('Certification email sent to: ') + course_participant)
                i +=1



#buddycert starts here
        elif course_certification == 'buddycert':

#Make a certificate for the course_participant
                doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\LTSBuddy.docx')
                obj_styles = doc.styles
                obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.PARAGRAPH)
                obj_font = obj_charstyle.font
                obj_font.size = Pt(80)
                obj_font.name = 'Freestyle Script'

#Update course_participant
                for paragraph in doc.paragraphs:
                        if 'name_here' in paragraph.text:
                                paragraph.style = doc.styles['CommentsStyle']
                                paragraph.text = course_participant
        
#Save, convert to pdf, delete .docx
                doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Buddy, ') + course_participant + ('.docx'))
                convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Buddy, ') + course_participant + ('.docx'))
                os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Buddy, ') + course_participant + ('.docx'))

#Outlook mail portion
                outlook = win32com.client.Dispatch('outlook.application')
                mail = outlook.CreateItem(0)
                mail.To = participant_email
                mail.Subject = 'Yachting New Zealand - Coaching Course Certificate - ' + course_participant 
                mail.HTMLBody = ("Dear ") + course_participant + (",<br><br>You have successfully completed the Yachting New Zealand Learn to Sail Coach Course. I am pleased to advise that you are now an officially recognised <b> Buddy Learn to Sail Coach. </b>"
                        "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                        "As a Buddy Learn to Sail Coach you are qualified to assist with the Yachting New Zealand Learn to Sail Dinghy (Level I and II) program. These levels can be taught at yacht clubs and other Yachting New Zealand affiliated organisations subject to the Yachting New Zealand safety requirements. <br><br>"
                        "A buddy coach must always be working under the supervision of a fully qualified Learn to Sail coach. <br><br>"
                        "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b>Buddy Learn to Sail Coach (Dinghy) Certificate.</b> You can upgrade to an Assistant Learn to sail coach when you turn 15 or working with the feedback the coach developer has given to upgrade your certificate. When you are ready to upgrade, you can by filling out a revalidation from, available to download off the Yachting New Zealand website.   <br><br>"
                        "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                        "If you are interested in furthering your coaching experience you may be keen to look at some becoming a Race Coach – information can be found under the <b>Coaches</b> page on the Yachting New Zealand website. Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                        "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                        "Regards,<br><br>") + email_signature

#Attachments
                course_certificate =  ('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Buddy, ') + course_participant + ('.pdf')
                mail.attachments.Add(course_certificate)
                mail.attachments.Add('c:\\Users\Peters\Documents\CDM\Certificates\Attachments\YACHTING NEW ZEALAND SAFETY REQUIREMENT.pdf')
                mail.Send()

                print(('Certification email sent to: ') + course_participant)
                i +=1



#keelboat 1 starts here
        if course_certification == 'k1':

#Make a certificate for the course_participant
                doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\K1.docx')
                obj_styles = doc.styles
                obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.PARAGRAPH)
                obj_font = obj_charstyle.font
                obj_font.size = Pt(80)
                obj_font.name = 'Freestyle Script'

#Update course_participant
                for paragraph in doc.paragraphs:
                        if 'name_here' in paragraph.text:
                                paragraph.style = doc.styles['CommentsStyle']
                                paragraph.text = course_participant
        
#Save, convert to pdf, delete .docx
                doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat 1, ') + course_participant + ('.docx'))
                convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat 1, ') + course_participant + ('.docx'))
                os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat 1, ') + course_participant + ('.docx'))

#Outlook mail portion
                outlook = win32com.client.Dispatch('outlook.application')
                mail = outlook.CreateItem(0)
                mail.To = participant_email
                mail.Subject = 'Yachting New Zealand - Coaching Course Certificate - ' + course_participant 
                mail.HTMLBody = KBbody

                course_certificate =  ('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat 1, ') + course_participant + ('.pdf')
                mail.attachments.Add(course_certificate)
                mail.Send()

                print(('Certification email sent to: ') + course_participant)
                i +=1



#keelboat 2 starts here
        if course_certification == 'k2':

#Make a certificate for the course_participant
                doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\K2.docx')
                obj_styles = doc.styles
                obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.PARAGRAPH)
                obj_font = obj_charstyle.font
                obj_font.size = Pt(80)
                obj_font.name = 'Freestyle Script'

#Update course_participant
                for paragraph in doc.paragraphs:
                        if 'name_here' in paragraph.text:
                                paragraph.style = doc.styles['CommentsStyle']
                                paragraph.text = course_participant
        
#Save, convert to pdf, delete .docx
                doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat 2, ') + course_participant + ('.docx'))
                convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat 2, ') + course_participant + ('.docx'))
                os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat 2, ') + course_participant + ('.docx'))

#Outlook mail portion
                outlook = win32com.client.Dispatch('outlook.application')
                mail = outlook.CreateItem(0)
                mail.To = participant_email
                mail.Subject = 'Yachting New Zealand - Coaching Course Certificate - ' + course_participant 
                mail.HTMLBody = KBbody

                course_certificate =  ('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat 2, ') + course_participant + ('.pdf')
                mail.attachments.Add(course_certificate)
                mail.Send()

                print(('Certification email sent to: ') + course_participant)
                i +=1



#keelboat 3 starts here
        if course_certification == 'k3':

#Make a certificate for the course_participant
                doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\K3.docx')
                obj_styles = doc.styles
                obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.PARAGRAPH)
                obj_font = obj_charstyle.font
                obj_font.size = Pt(80)
                obj_font.name = 'Freestyle Script'

#Update course_participant
                for paragraph in doc.paragraphs:
                        if 'name_here' in paragraph.text:
                                paragraph.style = doc.styles['CommentsStyle']
                                paragraph.text = course_participant
        
#Save, convert to pdf, delete .docx
                doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat 3, ') + course_participant + ('.docx'))
                convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat 3, ') + course_participant + ('.docx'))
                os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat 3, ') + course_participant + ('.docx'))

#Outlook mail portion
                outlook = win32com.client.Dispatch('outlook.application')
                mail = outlook.CreateItem(0)
                mail.To = participant_email
                mail.Subject = 'Yachting New Zealand - Coaching Course Certificate - ' + course_participant 
                mail.HTMLBody = KBbody

                course_certificate =  ('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat 3, ') + course_participant + ('.pdf')
                mail.attachments.Add(course_certificate)
                mail.Send()

                print(('Certification email sent to: ') + course_participant)
                i +=1

#You need to complete embark
        if course_certification == 'doembark':
                outlook = win32com.client.Dispatch('outlook.application')
                mail = outlook.CreateItem(0)
                mail.To = participant_email
                mail.Subject = 'Yachting New Zealand - Coaching Course Certificate - ' + course_participant 
                mail.HTMLBody = ("Dear ") + course_participant + (",<br><br>The Yachting New Zealand Coach Course you were on has finished, but have yet to complete the EMBARK online learning portion of the course. The section of the course that you still need to complete is either part of, or all of, the <b>'Coach Yachting 101'</b> section. There are 4 modules: <br> "
                        " <b>Embark Introduction </b><br><b>Coach Code of Conduct </b><br><b>Safety Requirements </b><br><b> Quality Coaching </b><br><br>"
                        "If you are having trouble accessing EMBARK, I have attached a form that explains how to access the EMBARK learning. <br><br>"
                        "Once you have finished, please be sure to email me that you have completed, and I can update the Yachting New Zealand records and get your certificate to you. <br><br>"
                        "Alternatively, if you are having trouble accessing EMBARK, please reach out after having tried to login.<br><br>"
                        "Regards,<br><br>") + email_signature

                mail.attachments.Add('c:\\Users\Peters\Documents\CDM\Certificates\Attachments\Accessing Embark.pdf')

                mail.Send()
                print(('YOU NEED TO DO EMBARK! >>>   ') + course_participant + ('   <<<'))
                i +=1 

print(" ")
print("DAYUMN, are we already done?!")
print("That happened WAY too fast. Maybe you should take the rest of the day off.")