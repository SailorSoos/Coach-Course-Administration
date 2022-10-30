import pandas as pd
import os
import docx
import win32com.client
import os.path
from docx.api import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx2pdf import convert

import time

excel_sheet = pd.read_excel("c:/Users/Peters/Downloads/process_these.xlsx")
registrations_expiry = '30 June 2026'

#universals - to separate to .common.py local import
#includes certreval and postprocess common
email_signature = ('<b><font color="rgb(0,65,92)"> Peter Soosalu | Coach Development Manager | Yachting New Zealand </font></b> <br>'
    '<b><font color="rgb(0,65,92)">M</b></font> <font color="rgb(0,65,92)">(021) 037 2419 </font>| <b><font color="rgb(0,65,92)">E</font></b> peters@yachting.org.nz <br>'
    '<a href="http://www.yachtingnz.org.nz">Yachtingnz.org.nz</a> | <a href="https://www.facebook.com/YachtingNewZealand/">Facebook</a> | <a href="https://www.facebook.com/NZLSailingTeam/">NZL Sailing Team</a>'
    '<br><br> For the latest news and offerings download the Yachting New Zealand app.'
    '<br><a href="https://apps.apple.com/us/app/yachting-nz-app/id1040333130?amp%3Bls=1&amp%3Bamp%3Bmt=8&amp%3Bl=nl"><img src="https://developer.apple.com/news/images/download-on-the-app-store-badge.png" width="108" height="36" alt="app_store"></a> <a href="https://play.google.com/store/apps/details?id=com.app.p1107GC"><img src="https://encrypted-tbn0.gstatic.com/images?q=tbn:ANd9GcT_E6ZM5CF_cPm4tzqW6MpFGm2efBY1QL6v6w&usqp=CAU" width="108" height="36" alt="android_store"></a>'
    '<br><br><img src="https://i.ibb.co/w7Lfz35/99798ce7-9981-4308-8295-3b3fa5386314.jpg" alt="ynz_banner">'
    '<br><br> <font size="-2" color="rgb(220,220,220)"> The content of this e-mail is confidential and may contain copyright information. If you are not the intended recipient, please delete the </font><br> <font size="-2" color="rgb(220,220,220)">message and notify the sender immediately. You should scan this message and any attached files for viruses. We accept no liability for </font><br><font size="-2" color="rgb(220,220,220)"> any loss caused either directly or indirectly by a virus arising from the use of this message or any attached file. Thank you. </font>')


def certificate(course_participant, course_type, template_path, output_path):
    if course_type == 'doembark':
        return
    doc = docx.Document(template_path + course_type + '.docx')
    obj_styles = doc.styles
    obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.PARAGRAPH)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(80)
    obj_font.name = 'Freestyle Script'

    for paragraph in doc.paragraphs:
            if 'name_here' in paragraph.text:
                    paragraph.style = doc.styles['CommentsStyle']
                    paragraph.text = course_participant

    doc.save((output_path + '\Yachting New Zealand - ') + course_participant + ('.docx'))
    convert((output_path + '\Yachting New Zealand - ') + course_participant + ('.docx'))
    os.remove((output_path + '\Yachting New Zealand - ') + course_participant + ('.docx'))

    time.sleep(1)

def email(course_participant, course_certificate, participant_email, course_type):
    if course_type == 'doembark':
        body = (("Dear ") + course_participant + (",<br><br>The Yachting New Zealand Coach Course you were on has finished, but have yet to complete the EMBARK online learning portion of the course. The section of the course that you still need to complete is either part of, or all of, the <b>'Coach Yachting 101'</b> section. There are 4 modules: <br> "
                " <b>Embark Introduction </b><br><b>Coach Code of Conduct </b><br><b>Safety Requirements </b><br><b> Quality Coaching </b><br><br>"
                "If you are having trouble accessing EMBARK, I have attached a form that explains how to access the EMBARK learning. <br><br>"
                "Once you have finished, please be sure to email me that you have completed, and I can update the Yachting New Zealand records and get your certificate to you. <br><br>"
                "Alternatively, if you are having trouble accessing EMBARK, please reach out after having tried to login.<br><br>"
                "Regards,<br><br>") + email_signature)
    elif course_type == 'Learn to Sail' or 'Learn to Sail Assistant' or 'Learn to Sail Buddy' or 'Learn to Sail Head':
        if course_type == 'Learn to Sail':
            courseText = "As a Learn to Sail coach you are qualified to teach all aspects of the Yachting New Zealand Learn to Sail Dinghy (Level I and II) program. These levels can be taught at yacht clubs and other Yachting New Zealand affiliated organisations subject to the Yachting New Zealand safety requirements. <br><br>"
        elif course_type == "Learn to Sail Head":
            courseText = "As a Learn to Sail Head coach you are qualified to teach all aspects of the Yachting New Zealand Learn to Sail Dinghy (Level I and II) program. These levels can be taught at yacht clubs and other Yachting New Zealand affiliated organisations subject to the Yachting New Zealand safety requirements. <br><br>"
        elif course_type == "Learn to Sail Assistant":
            courseText = "As an assistant Learn to Sail coach you are qualified to assist with the Yachting New Zealand Learn to Sail Dinghy (Level I and II) program. These levels can be taught at yacht clubs and other Yachting New Zealand affiliated organisations subject to the Yachting New Zealand safety requirements. <br><br>"
        elif course_type == "Learn to Sail Buddy":
            courseText = "As an Buddy Learn to Sail coach you are qualified to assist with the Yachting New Zealand Learn to Sail Dinghy (Level I and II) program. These levels can be taught at yacht clubs and other Yachting New Zealand affiliated organisations subject to the Yachting New Zealand safety requirements. <br><br>"
        elif course_type == 'Keelboat 1' or 'Keelboat 2' or 'Keelboat 3':
            courseText = "As a keelboat coach, mentoring and supporting other coaches in your area is not only encouraged, but can help improve your own coaching experience by sharing ideas and seeing new ones. <br><br> Your qualification does not have a set required sailor to coach ratio, but it is recommended to have a max ratio of 7:1.<br><br>"
        body = (("Dear ") + course_participant + (",<br><br>You have successfully completed the Yachting New Zealand ") + course_type +(" Coach Course. I am pleased to advise that you are now an officially recognised <b> ") + course_type + (" Coach. </b>"
                "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>")
                 + courseText + ("Your qualification is valid until ") + registrations_expiry + (" at which time you will need to revalidate. Please find enclosed your <b>") + course_type + ("</b> certificate. You can upgrade to a Learn to sail coach when you turn 18 or working with the feedback the coach developer has given to upgrade your certificate. When you are ready to upgrade, you can by filling out a revalidation from, available to download off the Yachting New Zealand website.   <br><br>"
                "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                "If you are interested in furthering your coaching experience you may be keen to look at some becoming a Race Coach – information can be found under the <b>Coaches</b> page on the Yachting New Zealand website. Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                "Regards,<br><br>") + email_signature)
    elif course_type == 'Race':
        body = (("Dear ") + course_participant + (",<br><br>I am pleased to advise that you are now an officially recognised <b> Race Coach. </b>"
                "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                "The Race Coach is the first step along the race coach pathway, followed by the Regatta coach, Performance coach, and finally Olympic coach. As a race coach, mentoring and supporting other coaches in your area is not only encouraged, but can help improve your own coaching experience by sharing ideas and seeing new ones. <br><br>"
                "Your qualification does not have a set required sailor to coach ratio, but it is recommended to have a max ratio of 6:1.<br><br>"
                "Your qualification is valid until ") + registrations_expiry + (" at which time you will need to revalidate. Please find enclosed your <b> Race Coach Certificate.</b> <br><br>"
                "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                "Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                "Regards,<br><br>") + email_signature)
    
    outlook = win32com.client.Dispatch('outlook.application')
    mail = outlook.CreateItem(0)
    mail.To = participant_email
    mail.Subject = 'Yachting New Zealand - Coaching Course Certificate - ' + course_participant 
    mail.HTMLBody = body
    course_certificate = ('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - ') + course_participant + ('.pdf')
    if course_type == 'doembark':
        mail.attachments.Add('c:\\Users\Peters\Documents\CDM\Certificates\Attachments\Accessing Embark.pdf')
    else:
        mail.attachments.Add(course_certificate)
        mail.attachments.Add('c:\\Users\Peters\Documents\CDM\Certificates\Attachments\YACHTING NEW ZEALAND SAFETY REQUIREMENT.pdf')

    mail.Send()
    print("Email sent to: " + course_participant)

i = 0
while i < len(excel_sheet):
    registrations_name1 = excel_sheet.iloc[i, 3]
    registrations_name = str(registrations_name1)
    registrations_qual1 = excel_sheet.iloc[i, 7]
    registrations_qual = str(registrations_qual1)
    registrations_email1 = excel_sheet.iloc[i, 11]
    registrations_email = str(registrations_email1)

    if registrations_qual == 'Assistant LTS Coach (Dinghy)':
            course_type = 'Learn to Sail Assistant'
    elif registrations_qual == 'LTS Coach (Dinghy)':
            course_type = 'Learn to Sail'
    elif registrations_qual == 'LTS Buddy':
            course_type = 'Learn to Sail Buddy'
    elif registrations_qual == 'Master LTS Coach (Dinghy)':
            course_type = 'Learn to Sail Head'
    elif registrations_qual == 'LTS Coach (Keelboat) Level 1':
            course_type = 'Keelboat 1'
    elif registrations_qual == 'LTS Coach (Keelboat) Level 2':
            course_type = 'Keelboat 2'
    elif registrations_qual == 'LTS Coach (Keelboat) Level 3':
            course_type = 'Keelboat 3'
    elif registrations_qual == 'Race Coach':
            course_type = 'Race'
    else:
            course_type = 'doembark'
    
    template_path = ('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\\')
    output_path = 'c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs'

    certificate(registrations_name, course_type, template_path, output_path)
    email(registrations_name, course_type, registrations_email, course_type)
    i +=1

print("\n DAYUMN, are we already done?! \n That happened WAY too fast. Maybe you should take the rest of the day off. \n")






#this is the certreval script
course_participant = input("Participant name:")
print("")
participant_email = input("Email:")
print("\nCert type from list below:")
print("'assistcert' \n'ltscert' \n'headcert' \n'racecert' \n'keel1cert' \n'keel2cert' \n'keel2cert' \n'keel3cert' \n")
certification_type = input("Cert type: ")
#Update this date yearly
expiry_date = "30 June, 2025"

reval_cert_type = "Learn to Sail Coach"

#added from the postprocess concept
if course_type == 'Learn to Sail':
    courseText = "As a Learn to Sail coach you are qualified to teach all aspects of the Yachting New Zealand Learn to Sail Dinghy (Level I and II) program. These levels can be taught at yacht clubs and other Yachting New Zealand affiliated organisations subject to the Yachting New Zealand safety requirements. <br><br>"
elif course_type == "Learn to Sail Head":
    courseText = "As a Learn to Sail Head coach you are qualified to teach all aspects of the Yachting New Zealand Learn to Sail Dinghy (Level I and II) program. These levels can be taught at yacht clubs and other Yachting New Zealand affiliated organisations subject to the Yachting New Zealand safety requirements. <br><br>"
elif course_type == "Learn to Sail Assistant":
    courseText = "As an assistant Learn to Sail coach you are qualified to assist with the Yachting New Zealand Learn to Sail Dinghy (Level I and II) program. These levels can be taught at yacht clubs and other Yachting New Zealand affiliated organisations subject to the Yachting New Zealand safety requirements. <br><br>"
elif course_type == "Learn to Sail Buddy":
    courseText = "As an Buddy Learn to Sail coach you are qualified to assist with the Yachting New Zealand Learn to Sail Dinghy (Level I and II) program. These levels can be taught at yacht clubs and other Yachting New Zealand affiliated organisations subject to the Yachting New Zealand safety requirements. <br><br>"
elif course_type == 'Keelboat 1' or 'Keelboat 2' or 'Keelboat 3':
    courseText = "As a keelboat coach, mentoring and supporting other coaches in your area is not only encouraged, but can help improve your own coaching experience by sharing ideas and seeing new ones. <br><br> Your qualification does not have a set required sailor to coach ratio, but it is recommended to have a max ratio of 7:1.<br><br>"


#Need to customize email for the qualification
if certification_type == 'assistcert':
        email = ("Dear ") + course_participant + (",<br><br>After reviewing your revalidation request, I am happy to revalidate your Yachting New Zealand") + reval_cert_type + ("certificate. You are now an officially recognised <b> Assistant Learn to Sail Coach. </b>"
                        "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                        "As an Assistant Learn to Sail Coach you are qualified to assist with the Yachting New Zealand Learn to Sail Dinghy (Level I and II) program. These levels can be taught at yacht clubs and other Yachting New Zealand affiliated organisations subject to the Yachting New Zealand safety requirements. <br><br>"
                        "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b>Assistant Learn to Sail Coach (Dinghy) Certificate.</b> You can upgrade to a Learn to sail coach when you turn 18 or working with the feedback the coach developer has given to upgrade your certificate. When you are ready to upgrade, you can by filling out a revalidation from, available to download off the Yachting New Zealand website.   <br><br>"
                        "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                        "If you are interested in furthering your coaching experience you may be keen to look at some becoming a Race Coach – information can be found under the <b>Coaches</b> page on the Yachting New Zealand website. Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                        "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                        "Regards,<br><br>") + email_signature
elif certification_type == 'ltscert':
        email = ("Dear ") + course_participant + (",<br><br>After reviewing your revalidation request, I am happy to revalidate your Yachting New Zealand Learn to Sail coach certificate. You are now an officially recognised <b> Learn to Sail Coach. </b>"
                        "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                        "As a Learn to Sail Coach you are qualified to teach all aspects of the Yachting New Zealand Learn to Sail Dinghy (Level I and II) program. These levels can be taught at yacht clubs and other Yachting New Zealand affiliated organisations subject to the Yachting New Zealand safety requirements. <br><br>"
                        "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b> Learn to Sail Coach (Dinghy) Certificate.</b> You can upgrade to a Learn to sail coach when you turn 18 or working with the feedback the coach developer has given to upgrade your certificate. When you are ready to upgrade, you can by filling out a revalidation from, available to download off the Yachting New Zealand website.   <br><br>"
                        "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                        "If you are interested in furthering your coaching experience you may be keen to look at some becoming a Race Coach – information can be found under the <b>Coaches</b> page on the Yachting New Zealand website. Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                        "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                        "Regards,<br><br>") + email_signature
elif certification_type == 'headcert':
        email = ("Dear ") + course_participant + (",<br><br>After reviewing your revalidation request, I am happy to revalidate your Yachting New Zealand Learn to Sail coach certificate. You are now an officially recognised <b> Head Learn to Sail Coach. </b>"
                        "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                        "As a Head Learn to Sail Coach you are qualified to teach all aspects of the Yachting New Zealand Learn to Sail Dinghy (Level I and II) program. These levels can be taught at yacht clubs and other Yachting New Zealand affiliated organisations subject to the Yachting New Zealand safety requirements. Additionally, as a Head Learn to Sail coach, mentoring and supporting other coaches in your area is not only encouraged, but can help improve your own coaching experience by sharing ideas and seeing new ones. <br><br>"
                        "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b> Head Learn to Sail Coach (Dinghy) Certificate.</b> <br><br>"
                        "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                        "If you are interested in furthering your coaching experience you may be keen to look at some becoming a Race Coach – information can be found under the <b>Coaches</b> page on the Yachting New Zealand website. Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                        "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                        "Regards,<br><br>") + email_signature
elif certification_type == 'racecert':
        email = ("Dear ") + course_participant + (",<br><br>After reviewing your revalidation request, I am happy to revalidate your Yachting New Zealand Race coach certificate. You are now an officially recognised <b> Race Coach. </b>"
                        "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                        "Additionally, as a Race coach, mentoring and supporting other coaches in your area is not only encouraged, but can help improve your own coaching experience by sharing ideas and seeing new ones. <br><br>"
                        "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b> Race Coach Certificate.</b> <br><br>"
                        "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                        "If you are interested in furthering your coaching experience you may be keen to look at some becoming a Regatta Coach – information can be found under the <b>Coaches</b> page on the Yachting New Zealand website. Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                        "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                        "Regards,<br><br>") + email_signature
elif certification_type == 'keel1cert':
        email = ("Dear ") + course_participant + (",<br><br>After reviewing your revalidation request, I am happy to revalidate your Yachting New Zealand Keelboat Level 1 coach certificate. You are now an officially recognised <b> Level 1 Keelboat Coach. </b>"
                        "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                        "Additionally, as a Keelboat coach, mentoring and supporting other coaches in your area is not only encouraged, but can help improve your own coaching experience by sharing ideas and seeing new ones. <br><br>"
                        "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b> Keelboat Level 1 Coach Certificate.</b> <br><br>"
                        "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                        "Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                        "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                        "Regards,<br><br>") + email_signature
elif certification_type == 'keel2cert':
        email = ("Dear ") + course_participant + (",<br><br>After reviewing your revalidation request, I am happy to revalidate your Yachting New Zealand Keelboat Level 2 coach certificate. You are now an officially recognised <b> Level 2 Keelboat Coach. </b>"
                        "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                        "Additionally, as a Keelboat coach, mentoring and supporting other coaches in your area is not only encouraged, but can help improve your own coaching experience by sharing ideas and seeing new ones. <br><br>"
                        "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b> Keelboat Level 2 Coach Certificate.</b> <br><br>"
                        "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                        "Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                        "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                        "Regards,<br><br>") + email_signature
elif certification_type == 'keel3cert':
        email = ("Dear ") + course_participant + (",<br><br>After reviewing your revalidation request, I am happy to revalidate your Yachting New Zealand Keelboat Level 3 coach certificate. You are now an officially recognised <b> Level 3 Keelboat Coach. </b>"
                        "As part of an effort to reduce the carbon footprint certificates make, I have attached your certificate as a pdf for you to download. If you would like a hard copy of the certificate, please let me know, along with your mailing address, and I can make sure it gets to the right place. <br><br>"
                        "Additionally, as a Keelboat coach, mentoring and supporting other coaches in your area is not only encouraged, but can help improve your own coaching experience by sharing ideas and seeing new ones. <br><br>"
                        "Your qualification is valid until ") + expiry_date + (" at which time you will need to revalidate. Please find enclosed your <b> Keelboat Level 3 Coach Certificate.</b> <br><br>"
                        "Although not mandatory, we do highly recommend that the following courses be added to your qualifications: RYA Powerboat level 2, current first aid certificate and you may also consider a VHF Operators Certificate. Information can be found on the Yachting New Zealand and Coastguard Boating Education websites.<br><br>"
                        "Also check out the Yachting New Zealand Coaches Forum Facebook group.<br><br>"
                        "Congratulations again on attaining this qualification. Please do not hesitate to contact me at any time if you require any additional assistance or guidance during your coaching.<br><br>"
                        "Regards,<br><br>") + email_signature

#Open the certificate template
if certification_type == 'assistcert':
        doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\LTSAssistant.docx')
elif certification_type == 'ltscert':
        doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\LTS.docx')
elif certification_type == 'buddycert':
        doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\LTSBuddy.docx')
elif certification_type == 'headcert':
        doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\LTSHead.docx')
elif certification_type == 'racecert':
        doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\Race.docx')
elif certification_type == 'keel1cert':
        doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\K1.docx')
elif certification_type == 'keel2cert':
        doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\K2.docx')
elif certification_type == 'keel3cert':
        doc = docx.Document('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Templates\K3.docx')

#Change the font size and type
obj_styles = doc.styles
obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.PARAGRAPH)
obj_font = obj_charstyle.font
obj_font.size = Pt(80)
obj_font.name = 'Freestyle Script'

#Update course_participant
for paragraph in doc.paragraphs:
        if 'name_here' in paragraph.text:
                paragraph.style = doc.styles['CommentsStyle']
                paragraph.text = course_participant
        
#Where to save it, what to save it as in .docx
if certification_type == 'assistcert':
        doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Assistant, ') + course_participant + ('.docx'))
elif certification_type == 'ltscert':
        doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS, ') + course_participant + ('.docx'))
elif certification_type == 'headcert':
        doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Head Coach, ') + course_participant + ('.docx'))
elif certification_type == 'racecert':
        doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Race Coach, ') + course_participant + ('.docx'))
elif certification_type == 'keel1cert':
        doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 1, ') + course_participant + ('.docx'))
elif certification_type == 'keel2cert':
        doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 2, ') + course_participant + ('.docx'))
elif certification_type == 'keel3cert':
        doc.save(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 3, ') + course_participant + ('.docx'))

#Convert word doc to pdf
from docx2pdf import convert
if certification_type == 'assistcert':
        convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Assistant, ') + course_participant + ('.docx'))
elif certification_type == 'ltscert':
        convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS, ') + course_participant + ('.docx'))
elif certification_type == 'headcert':
        convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Head Coach, ') + course_participant + ('.docx'))
elif certification_type == 'racecert':
        convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Race Coach, ') + course_participant + ('.docx'))
elif certification_type == 'keel1cert':
        convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 1, ') + course_participant + ('.docx'))
elif certification_type == 'keel2cert':
        convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 2, ') + course_participant + ('.docx'))
elif certification_type == 'keel3cert':
        convert(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 3, ') + course_participant + ('.docx'))

#Delete the word document version
if certification_type == 'assistcert':
        os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Assistant, ') + course_participant + ('.docx'))
elif certification_type == 'ltscert':
        os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS, ') + course_participant + ('.docx'))
elif certification_type == 'headcert':
        os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Head Coach, ') + course_participant + ('.docx'))
elif certification_type == 'racecert':
        os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Race Coach, ') + course_participant + ('.docx'))
elif certification_type == 'keel1cert':
        os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 1, ') + course_participant + ('.docx'))
elif certification_type == 'keel2cert':
        os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 2, ') + course_participant + ('.docx'))
elif certification_type == 'keel3cert':
        os.remove(('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 3, ') + course_participant + ('.docx'))


#Outlook mail portion
import win32com.client
import os.path
outlook = win32com.client.Dispatch('outlook.application')
mail = outlook.CreateItem(0)

#Create the email 
mail.To = participant_email
mail.Subject = 'Yachting New Zealand - Coaching Course Certificate - ' + course_participant 
mail.HTMLBody = email

# Certification attachments
if certification_type == 'assistcert':
        course_certificate =  (('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Assistant, ') + course_participant + ('.pdf'))
elif certification_type == 'ltscert':
        course_certificate =  (('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS, ') + course_participant + ('.pdf'))
elif certification_type == 'headcert':
        course_certificate =  (('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - LTS Head Coach, ') + course_participant + ('.pdf'))
elif certification_type == 'racecert':
        course_certificate =  (('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Race Coach, ') + course_participant + ('.pdf'))
elif certification_type == 'keel1cert':
        course_certificate =  (('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 1, ') + course_participant + ('.pdf'))
elif certification_type == 'keel2cert':
        course_certificate =  (('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 2, ') + course_participant + ('.pdf'))
elif certification_type == 'keel3cert':
        course_certificate =  (('c:\\Users\Peters\Documents\CDM\Certificates\Auto_created_certs\Yachting New Zealand - Keelboat Level 3, ') + course_participant + ('.pdf'))
mail.attachments.Add(course_certificate)
mail.attachments.Add('c:\\Users\Peters\Documents\CDM\Certificates\Attachments\YACHTING NEW ZEALAND SAFETY REQUIREMENT.pdf')
mail.Send()

print("")
print(('Revalidation email sent to: ') + course_participant)

#cc another person (alt email from crm?)
#mail.CC = 'peters@yachtingnz.org.nz'